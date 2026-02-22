"""Generate MedSift AI project writeup as a DOCX file."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# --- Styles ---
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

# --- Title Page ---
for _ in range(6):
    doc.add_paragraph("")

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("MedSift AI")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0, 102, 153)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Privacy-First AI Medical Scribe & Clinical Decision Support")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph("")

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run("GTHack 2026")
run.font.size = Pt(14)
run.bold = True

doc.add_paragraph("")

team_info = doc.add_paragraph()
team_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = team_info.add_run("Project Writeup")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# ===== TABLE OF CONTENTS =====
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Introduction",
    "2. Problem Statement",
    "3. Solution Overview",
    "4. System Architecture",
    "5. Technology Stack",
    "6. Core Features",
    "7. Privacy & HIPAA Compliance",
    "8. API Design",
    "9. User Interface",
    "10. How It Works (End-to-End Pipeline)",
    "11. External Integrations",
    "12. Testing & Validation",
    "13. Future Scope",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ===== 1. INTRODUCTION =====
doc.add_heading("1. Introduction", level=1)
doc.add_paragraph(
    "MedSift AI is a privacy-first, fully local AI-powered medical scribe and clinical decision support system. "
    "It transforms doctor-patient conversations into structured medical documentation, actionable clinical insights, "
    "and patient-friendly after-visit summaries — all while ensuring HIPAA compliance through complete local processing "
    "and automated Protected Health Information (PHI) redaction."
)
doc.add_paragraph(
    "Unlike cloud-based medical transcription services, MedSift AI processes everything on the clinician's own machine. "
    "No audio, no transcripts, and no patient data ever leaves the local environment. This eliminates the privacy risks "
    "associated with third-party cloud services while delivering the same AI-powered capabilities."
)

# ===== 2. PROBLEM STATEMENT =====
doc.add_heading("2. Problem Statement", level=1)
doc.add_paragraph(
    "Clinicians spend an estimated 2 hours on documentation for every 1 hour of patient care. This administrative burden "
    "leads to physician burnout, reduced patient interaction time, and documentation errors. Existing solutions often:"
)
bullets = [
    "Send sensitive patient data to cloud services, creating HIPAA compliance risks",
    "Produce unstructured notes that require manual reformatting",
    "Lack integration with clinical research (trials, literature)",
    "Don't provide patient-friendly summaries for after-visit communication",
    "Offer no feedback mechanism for clinicians to improve extraction quality over time",
]
for b in bullets:
    doc.add_paragraph(b, style="List Bullet")

doc.add_paragraph(
    "MedSift AI addresses all of these challenges with a fully local, privacy-preserving system that automates "
    "the entire documentation workflow — from audio capture to PDF generation."
)

# ===== 3. SOLUTION OVERVIEW =====
doc.add_heading("3. Solution Overview", level=1)
doc.add_paragraph("MedSift AI provides an end-to-end pipeline that:")
steps = [
    ("Records & Transcribes", "Captures doctor-patient conversations and converts speech to text using OpenAI Whisper with pause-based speaker diarization."),
    ("Redacts PHI", "Automatically detects and masks Protected Health Information (names, SSNs, phone numbers, medical record numbers) using Microsoft Presidio before any data is processed or stored."),
    ("Extracts Structured Data", "Uses a local LLM (LLaMA 3.1 via Ollama) to extract a clinician-facing SOAP note and a patient-facing summary with medications, tests, follow-ups, and lifestyle recommendations."),
    ("Validates Extractions", "Cross-references every extracted item against the original transcript to verify accuracy, marking items as verified or unverified."),
    ("Searches Clinical Research", "Queries ClinicalTrials.gov and PubMed/Semantic Scholar to find relevant clinical trials and recent literature based on the patient's conditions and medications."),
    ("Generates PDFs", "Produces branded After Visit Summary PDFs that doctors can review, edit, and hand to patients."),
    ("Learns from Feedback", "Incorporates clinician feedback to boost search keywords and improve future literature recommendations."),
]
for title_text, desc in steps:
    p = doc.add_paragraph()
    run = p.add_run(f"{title_text}: ")
    run.bold = True
    p.add_run(desc)

# ===== 4. SYSTEM ARCHITECTURE =====
doc.add_heading("4. System Architecture", level=1)
doc.add_paragraph(
    "MedSift AI follows a modular, layered architecture with clear separation of concerns:"
)

doc.add_heading("Architecture Diagram", level=2)
arch_text = (
    "Audio Input\n"
    "    |\n"
    "    v\n"
    "OpenAI Whisper (Local) + Pause-Based Speaker Diarization\n"
    "    |\n"
    "    v\n"
    "Microsoft Presidio PHI Redaction (Local)\n"
    "    |\n"
    "    v\n"
    "Ollama LLaMA 3.1 (Local LLM)\n"
    "    |              |\n"
    "    v              v\n"
    "Patient Summary   Clinician SOAP Note\n"
    "    |              |\n"
    "    v              v\n"
    "Post-Extraction Validation (Evidence Grounding)\n"
    "    |\n"
    "    +---> ClinicalTrials.gov API\n"
    "    +---> PubMed / Semantic Scholar APIs\n"
    "    +---> PDF Generation (fpdf2)\n"
    "    |\n"
    "    v\n"
    "SQLite Database (FTS5 Full-Text Search)\n"
    "    |\n"
    "    v\n"
    "FastAPI REST Backend --> Streamlit Web UI"
)
p = doc.add_paragraph()
run = p.add_run(arch_text)
run.font.name = "Consolas"
run.font.size = Pt(9)

doc.add_heading("Layer Breakdown", level=2)
layers = [
    ("Presentation Layer", "Streamlit web application with multi-page navigation for upload, live transcription, visit history, and analytics."),
    ("API Layer", "FastAPI REST backend with 11+ endpoints handling transcription, analysis, CRUD operations, export, feedback, and analytics."),
    ("Core Processing Layer", "Modular Python services for transcription, PHI redaction, LLM extraction, validation, clinical trials search, literature search, and PDF generation."),
    ("Data Layer", "SQLite database with WAL mode, FTS5 full-text search, JSON storage for nested objects, and automated trigger-based indexing."),
    ("AI/ML Layer", "OpenAI Whisper for speech-to-text, pause-based speaker diarization, Microsoft Presidio + spaCy for NER-based PHI detection, and Ollama LLaMA 3.1 for structured extraction."),
]
for layer_name, layer_desc in layers:
    p = doc.add_paragraph()
    run = p.add_run(f"{layer_name}: ")
    run.bold = True
    p.add_run(layer_desc)

# ===== 5. TECHNOLOGY STACK =====
doc.add_heading("5. Technology Stack", level=1)

# Table
table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = "Category"
hdr[1].text = "Technology"
hdr[2].text = "Purpose"

tech_rows = [
    ("Backend Framework", "FastAPI + Uvicorn", "Async REST API server"),
    ("Data Validation", "Pydantic v2", "Schema validation & serialization"),
    ("Speech-to-Text", "OpenAI Whisper", "Local audio transcription (base model)"),
    ("Deep Learning", "PyTorch", "Backend for Whisper models"),
    ("PHI Redaction", "Microsoft Presidio", "Named entity recognition for HIPAA compliance"),
    ("NLP Engine", "spaCy (en_core_web_lg)", "NER pipeline for Presidio"),
    ("Local LLM", "Ollama + LLaMA 3.1 (8B)", "Structured data extraction from transcripts"),
    ("Database", "SQLite + FTS5", "Persistent storage with full-text search"),
    ("PDF Generation", "fpdf2", "After Visit Summary PDF export"),
    ("Web UI", "Streamlit", "Interactive demo interface"),
    ("HTTP Client", "Requests", "External API integration"),
    ("Configuration", "python-dotenv", "Environment variable management"),
    ("Testing", "pytest + httpx", "Unit and integration testing"),
    ("Audio Processing", "FFmpeg", "Audio format conversion"),
]
for cat, tech, purpose in tech_rows:
    row = table.add_row().cells
    row[0].text = cat
    row[1].text = tech
    row[2].text = purpose

# ===== 6. CORE FEATURES =====
doc.add_heading("6. Core Features", level=1)

doc.add_heading("6.1 Audio Transcription with Speaker Diarization", level=2)
doc.add_paragraph(
    "MedSift AI uses OpenAI's Whisper model running locally for speech-to-text conversion. "
    "It supports multiple audio formats (.mp3, .wav, .m4a, .webm, .flac, .ogg) and automatically "
    "identifies speakers using a pause-based diarization heuristic, labeling dialogue as "
    "\"Doctor\" and \"Patient\"."
)
doc.add_paragraph(
    "Live transcription is also supported, processing audio in real-time chunks with stateful speaker "
    "tracking across segments and configurable session timeouts."
)

doc.add_heading("6.2 Automated PHI Redaction", level=2)
doc.add_paragraph(
    "Before any transcript is processed by the LLM or stored in the database, it passes through Microsoft "
    "Presidio's NER-based redaction pipeline. The system detects and masks:"
)
phi_types = [
    "Person names (PERSON)",
    "Phone numbers (PHONE_NUMBER)",
    "Email addresses (EMAIL)",
    "Physical locations (LOCATION)",
    "Social Security Numbers (US_SSN)",
    "Medical Record Numbers (custom recognizer: MRN-XXXXXX patterns)",
    "Insurance IDs (custom recognizer: INS-XXXXXX, policy#, member ID patterns)",
    "Driver's licenses (US_DRIVER_LICENSE)",
    "Credit card numbers (CREDIT_CARD)",
]
for phi in phi_types:
    doc.add_paragraph(phi, style="List Bullet")
doc.add_paragraph(
    "Each entity is replaced with an indexed tag (e.g., [PERSON_1], [PHONE_2]) to preserve context "
    "while removing identifiable information. Date/time values are intentionally preserved as they carry "
    "clinical significance."
)

doc.add_heading("6.3 AI-Powered SOAP Note Extraction", level=2)
doc.add_paragraph(
    "Using Ollama with LLaMA 3.1 (8B parameters) running entirely locally, MedSift AI extracts a structured "
    "SOAP note from the redacted transcript:"
)
soap = [
    ("Subjective (S)", "Patient's reported symptoms, complaints, and history — presented as detailed clinical bullet points."),
    ("Objective (O)", "Organized into subsections: Vital Signs, Physical Examination, Mental State Examination, and Lab Results. Each subsection contains bullet-point findings."),
    ("Assessment (A)", "Clinical impressions, diagnoses, and differential diagnoses as bullet-point findings."),
    ("Plan (P)", "Treatment plans, medications prescribed, tests ordered, referrals, and follow-up instructions as bullet-point findings."),
]
for section, desc in soap:
    p = doc.add_paragraph()
    run = p.add_run(f"{section}: ")
    run.bold = True
    p.add_run(desc)
doc.add_paragraph(
    "The SOAP note is fully editable by the clinician after extraction, allowing them to add missing details "
    "(e.g., vitals not mentioned in the conversation). Edits persist to the database and are reflected in exported PDFs."
)

doc.add_heading("6.4 Patient-Friendly Summary", level=2)
doc.add_paragraph(
    "Alongside the clinical SOAP note, MedSift AI generates a warm, patient-facing summary that includes:"
)
patient_items = [
    "A personalized patient letter (Dear [Patient's Name]...) summarizing the visit in plain language",
    "Medications list with name, dose, frequency, duration, and plain-language instructions",
    "Tests ordered with preparation instructions and timelines",
    "Follow-up plan with actionable checklist items",
    "Lifestyle recommendations with detailed guidance",
    "Red flags — warning signs that should prompt urgent medical attention",
    "Questions & Answers from the visit",
]
for item in patient_items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("6.5 Evidence Validation", level=2)
doc.add_paragraph(
    "Every extracted item (medication, test, follow-up, etc.) includes an evidence field containing a direct "
    "quote from the transcript. The validation engine cross-references each item against the source transcript "
    "using two strategies:"
)
doc.add_paragraph("1. Exact substring matching for direct quotes", style="List Bullet")
doc.add_paragraph("2. Loose keyword matching with a 60%+ overlap threshold for paraphrased content", style="List Bullet")
doc.add_paragraph(
    "Items are marked as \"verified\" or \"unverified\", giving clinicians confidence in which extractions are "
    "directly grounded in the conversation versus inferred by the AI."
)

doc.add_heading("6.6 Clinical Research Integration", level=2)
doc.add_paragraph(
    "Based on the conditions and medications extracted from the SOAP note, MedSift AI automatically searches:"
)
doc.add_paragraph("ClinicalTrials.gov API v2 — Finds actively recruiting clinical trials relevant to the patient's conditions", style="List Bullet")
doc.add_paragraph("PubMed (NCBI E-utilities) — Searches the gold standard medical literature database for clinical trials, reviews, and meta-analyses", style="List Bullet")
doc.add_paragraph("Semantic Scholar API — Provides broader academic coverage with citation metrics for impact assessment", style="List Bullet")
doc.add_paragraph(
    "Results from PubMed and Semantic Scholar are deduplicated and merged, with PubMed prioritized for equal citation counts."
)

doc.add_heading("6.7 PDF Export (After Visit Summary)", level=2)
doc.add_paragraph(
    "MedSift AI generates branded PDF documents using fpdf2 that serve as After Visit Summaries. These include:"
)
pdf_items = [
    "MedSift AI header branding with visit date and type",
    "AI-generated disclaimer in red",
    "Patient letter in a warm, readable format",
    "Medications table with auto-truncation for overflow",
    "Tests ordered table",
    "Follow-up plan with checkbox items",
    "Lifestyle recommendations",
    "Red flag warnings highlighted in red",
    "Q&A section",
    "Optional SOAP note appendix (for clinician copies)",
    "Page numbers and professional footer",
]
for item in pdf_items:
    doc.add_paragraph(item, style="List Bullet")
doc.add_paragraph(
    "Clinicians can review and approve the summary before export, with the ability to edit the patient letter, "
    "toggle individual items for inclusion, and optionally include the SOAP note."
)

doc.add_heading("6.8 Clinician Feedback Loop", level=2)
doc.add_paragraph(
    "MedSift AI includes a feedback system that enables clinicians to rate extraction accuracy and literature "
    "relevance. Feedback types include:"
)
doc.add_paragraph("Extraction accuracy: correct / incorrect / missing", style="List Bullet")
doc.add_paragraph("Literature relevance: relevant / not relevant", style="List Bullet")
doc.add_paragraph(
    "Positive feedback on literature results extracts keywords from paper titles and boosts them in future "
    "searches, creating an adaptive learning loop that improves recommendation quality over time."
)

# ===== 7. PRIVACY & HIPAA =====
doc.add_heading("7. Privacy & HIPAA Compliance", level=1)
doc.add_paragraph("MedSift AI is designed from the ground up for HIPAA compliance:")

privacy_items = [
    ("100% Local Processing", "All AI models (Whisper, LLaMA 3.1, Presidio) run on the clinician's machine. No patient data is sent to any cloud service."),
    ("PHI Redaction Before Storage", "Transcripts are redacted before being processed by the LLM or saved to the database. Even if unredacted text is sent to the analysis endpoint, it is re-redacted as a safety net."),
    ("No Raw PHI in Database", "The database only stores redacted transcripts with indexed placeholder tags."),
    ("External APIs Receive Only De-identified Data", "ClinicalTrials.gov and PubMed/Semantic Scholar only receive condition names and drug names — never patient identifiers."),
    ("Custom Recognizers", "Beyond standard Presidio entities, custom recognizers detect Medical Record Numbers (MRN-XXXXXX) and Insurance IDs (INS-XXXXXX, policy#, member IDs)."),
    ("Date Preservation", "Dates are intentionally NOT redacted as they carry critical clinical significance for treatment timelines."),
]
for title_text, desc in privacy_items:
    p = doc.add_paragraph()
    run = p.add_run(f"{title_text}: ")
    run.bold = True
    p.add_run(desc)

# ===== 8. API DESIGN =====
doc.add_heading("8. API Design", level=1)
doc.add_paragraph("MedSift AI exposes a RESTful API via FastAPI with the following endpoints:")

api_table = doc.add_table(rows=1, cols=3)
api_table.style = "Light Grid Accent 1"
api_hdr = api_table.rows[0].cells
api_hdr[0].text = "Method"
api_hdr[1].text = "Endpoint"
api_hdr[2].text = "Description"

endpoints = [
    ("POST", "/api/transcribe", "Upload audio, transcribe with Whisper, redact PHI"),
    ("POST", "/api/analyze", "Full pipeline: extract summaries, search trials & literature"),
    ("GET", "/api/visits", "List visits with search, tag filter, and pagination"),
    ("GET", "/api/visits/{id}", "Get full visit details"),
    ("DELETE", "/api/visits/{id}", "Delete a visit"),
    ("PUT", "/api/visits/{id}/clinician-note", "Update editable SOAP note"),
    ("PUT", "/api/visits/{id}/patient-summary", "Update patient summary"),
    ("GET", "/api/export/{id}/pdf", "Download After Visit Summary PDF"),
    ("POST", "/api/export/reviewed/pdf", "Export doctor-reviewed PDF"),
    ("GET", "/api/trials/{id}", "Search clinical trials for a visit"),
    ("GET", "/api/literature/{id}", "Search literature for a visit"),
    ("POST", "/api/feedback", "Submit clinician feedback"),
    ("GET", "/api/feedback/{id}", "Get feedback for a visit"),
    ("GET", "/api/analytics", "Dashboard analytics"),
]
for method, endpoint, desc in endpoints:
    row = api_table.add_row().cells
    row[0].text = method
    row[1].text = endpoint
    row[2].text = desc

# ===== 9. USER INTERFACE =====
doc.add_heading("9. User Interface", level=1)
doc.add_paragraph(
    "The Streamlit-based web interface provides an intuitive experience across multiple pages:"
)

doc.add_heading("Upload & Process Page", level=2)
doc.add_paragraph(
    "Clinicians upload audio files or record directly. The page shows a progress pipeline through "
    "transcription, PHI redaction, AI analysis, and results. Results are displayed in tabs: Transcript, "
    "SOAP Note (editable), Care Plan (patient summary), and Review & Approve (for PDF export with doctor sign-off)."
)

doc.add_heading("Live Transcription Page", level=2)
doc.add_paragraph(
    "Real-time transcription using chunked audio processing with stateful speaker tracking. "
    "Supports browser-based microphone capture for live doctor-patient conversations."
)

doc.add_heading("Visit History Page", level=2)
doc.add_paragraph(
    "Searchable list of all past visits with full-text search, tag filtering, and expandable detail views. "
    "Each visit shows the SOAP note (editable), patient summary, clinical trials, and literature results. "
    "Doctors can review and approve summaries, then export PDFs."
)

doc.add_heading("Analytics Dashboard", level=2)
doc.add_paragraph(
    "Aggregated metrics including total visits, common conditions and medications, visits over time, "
    "extraction accuracy rates, and feedback-boosted keywords."
)

# ===== 10. HOW IT WORKS =====
doc.add_heading("10. How It Works (End-to-End Pipeline)", level=1)

pipeline_steps = [
    ("Step 1: Audio Capture", "The clinician uploads an audio recording of the patient visit or uses live transcription via the browser microphone."),
    ("Step 2: Speech-to-Text", "OpenAI Whisper (base model, running locally) converts the audio to text segments with timestamps. A pause-based diarization heuristic identifies which segments belong to the Doctor and which to the Patient."),
    ("Step 3: PHI Redaction", "Microsoft Presidio scans the transcript using spaCy's NER model and custom regex recognizers to identify and mask all Protected Health Information with indexed tags like [PERSON_1]."),
    ("Step 4: LLM Extraction", "The redacted transcript is sent to Ollama (LLaMA 3.1, running locally) with carefully crafted prompts. Two separate extractions run: one for the clinician SOAP note and one for the patient-friendly summary. The LLM returns structured JSON which is parsed with retry logic and validated against Pydantic schemas."),
    ("Step 5: Evidence Validation", "Each extracted item's evidence field is cross-referenced against the original transcript. Items are marked as verified (grounded in the conversation) or unverified (potentially inferred)."),
    ("Step 6: Clinical Research", "Conditions from the Assessment and medications from the Patient Summary are used to query ClinicalTrials.gov for recruiting trials and PubMed/Semantic Scholar for relevant literature."),
    ("Step 7: Storage", "The complete visit record (transcript, SOAP note, patient summary, trials, literature) is saved to SQLite with FTS5 indexing for fast search."),
    ("Step 8: Review & Export", "The clinician reviews the extracted data, edits the SOAP note to add missing details, approves the patient summary, and generates a branded PDF After Visit Summary."),
]
for step_title, step_desc in pipeline_steps:
    doc.add_heading(step_title, level=2)
    doc.add_paragraph(step_desc)

# ===== 11. EXTERNAL INTEGRATIONS =====
doc.add_heading("11. External Integrations", level=1)

doc.add_heading("ClinicalTrials.gov API v2", level=2)
doc.add_paragraph(
    "Searches for actively recruiting clinical trials matching the patient's conditions and medications. "
    "Returns trial NCT IDs, titles, status, conditions, interventions, locations, and direct URLs. "
    "Only RECRUITING trials are shown to ensure relevance."
)

doc.add_heading("PubMed (NCBI E-utilities)", level=2)
doc.add_paragraph(
    "Searches the National Library of Medicine's PubMed database — the gold standard for medical literature. "
    "Filters for clinical trials, reviews, and meta-analyses. Uses the free E-utilities API with a rate limit "
    "of 3 requests per second. No API key required."
)

doc.add_heading("Semantic Scholar API", level=2)
doc.add_paragraph(
    "Provides broader academic coverage beyond PubMed, with citation count metrics for assessing research impact. "
    "Results are deduplicated against PubMed results and merged, with PubMed prioritized for equal citation counts."
)

# ===== 12. TESTING & VALIDATION =====
doc.add_heading("12. Testing & Validation", level=1)
doc.add_paragraph("MedSift AI includes a comprehensive test suite with 31 tests:")

test_table = doc.add_table(rows=1, cols=3)
test_table.style = "Light Grid Accent 1"
test_hdr = test_table.rows[0].cells
test_hdr[0].text = "Test Module"
test_hdr[1].text = "Tests"
test_hdr[2].text = "Coverage"

tests = [
    ("test_phi_redaction.py", "8", "Presidio detection, custom recognizers (MRN, Insurance), indexed tags, edge cases"),
    ("test_risk_scoring.py", "10", "Rule-based risk score calculation"),
    ("test_extraction.py", "8", "Mocked Ollama responses, JSON parsing, Pydantic validation, retry logic"),
    ("test_transcription.py", "5", "Mocked Whisper model, speaker diarization, format validation"),
]
for module, count, coverage in tests:
    row = test_table.add_row().cells
    row[0].text = module
    row[1].text = count
    row[2].text = coverage

# ===== 13. FUTURE SCOPE =====
doc.add_heading("13. Future Scope", level=1)
future_items = [
    "Multi-language support for transcription and extraction",
    "Integration with Electronic Health Record (EHR) systems via FHIR/HL7 APIs",
    "Fine-tuned medical LLM for improved extraction accuracy",
    "Multi-patient visit support (beyond 2-speaker diarization)",
    "Mobile application for bedside documentation",
    "Encrypted database with role-based access control",
    "ICD-10 and CPT code auto-suggestion from SOAP notes",
    "Integration with pharmacy systems for e-prescribing",
]
for item in future_items:
    doc.add_paragraph(item, style="List Bullet")

# ===== Save =====
output_path = "/Users/tejas/Documents/GTHack/MedSift_AI_Writeup.docx"
doc.save(output_path)
print(f"Writeup saved to: {output_path}")
